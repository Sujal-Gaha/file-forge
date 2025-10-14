from pathlib import Path
from typing import Optional
import PyPDF2
from docx import Document


class DocumentProcessor:
    def __init__(self, input_path: Path, output_path: Optional[Path] = None) -> None:
        """
        Initialize DocumentProcessor

        Args:
            input_path: Path to input document
            output_path: Optional output path

        Returns:
            None
        """
        self.input_path = input_path
        self.output_path = output_path
        self.input_format = input_path.suffix.lower().replace(".", "")

    def _pdf_to_txt(self, output_path: Path) -> Path:
        """Convert PDF to plain text"""
        text_content = []

        with open(self.input_path, "rb") as pdf_file:
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text_content.append(page.extract_text())

        with open(output_path, "w", encoding="utf-8") as txt_file:
            txt_file.write("\n\n".join(text_content))

        return output_path

    def _docx_to_txt(self, output_path: Path) -> Path:
        """Convert DOCX to plain text"""
        doc = Document(str(self.input_path))

        text_content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)

        with open(output_path, "w", encoding="utf-8") as txt_file:
            txt_file.write("\n\n".join(text_content))

        return output_path

    def _txt_to_docx(self, output_path: Path) -> Path:
        """Convert plain text to DOCX"""
        doc = Document()

        with open(self.input_path, "r", encoding="utf-8") as txt_file:
            content = txt_file.read()

            paragraphs = content.split("\n\n")

            for para_text in paragraphs:
                if para_text.strip():
                    doc.add_paragraph(para_text.strip())

        doc.save(str(output_path))
        return output_path

    def convert_document(self, output_format: str) -> Path:
        """
        Convert document to a different format

        Args:
            output_format: Target format (pdf, txt, docx)

        Returns:
            Path to converted document
        """
        output_format = output_format.lower().replace(".", "")

        if self.output_path is None:
            self.output_path = self.input_path.with_suffix(f".{output_format}")

        if self.input_format == "pdf" and output_format == "txt":
            return self._pdf_to_txt(self.output_path)
        elif self.input_format == "docx" and output_format == "txt":
            return self._docx_to_txt(self.output_path)
        elif self.input_format == "txt" and output_format == "docx":
            return self._txt_to_docx(self.output_path)
        else:
            raise ValueError(
                f"Conversion from {self.input_format} to {output_format} is not supported yet. "
                f"Supported conversions: PDF→TXT, DOCX→TXT, TXT→DOCX"
            )

    def extract_pdf_pages(
        self, output_path: Path, start_page: int, end_page: Optional[int] = None
    ) -> Path:
        """
        Extract specific pages from a PDF

        Args:
            output_path: Path to output PDF
            start_page: Starting page number (1-indexed)
            end_page: Ending page number (1-indexed), None for last page

        Returns:
            Path to extracted PDF
        """
        with open(self.input_path, "rb") as pdf_file:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            pdf_writer = PyPDF2.PdfWriter()

            total_pages = len(pdf_reader.pages)

            start_idx = start_page - 1
            end_idx = (end_page if end_page else total_pages) - 1

            if start_idx < 0 or start_idx >= total_pages:
                raise ValueError(
                    f"Start page {start_page} is out of rang (1-{total_pages})"
                )
            if end_idx >= total_pages:
                raise ValueError(
                    f"End page {end_page} is out of range (1-{total_pages})"
                )

            for page_num in range(start_idx, end_idx + 1):
                pdf_writer.add_page(pdf_reader.pages[page_num])

            with open(output_path, "wb") as output_file:
                pdf_writer.write(output_file)

        return output_path

    def merge_pdfs(self, input_paths: list[Path], output_path: Path) -> Path:
        """
        Merge multiple PDFs into one

        Args:
            input_paths: List of PDF paths to merge
            output_path: Path to output merged PDF

        Returns:
            Path to merged PDF
        """
        pdf_write = PyPDF2.PdfWriter()

        for pdf_path in input_paths:
            with open(pdf_path, "rb") as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                for page in pdf_reader.pages:
                    pdf_write.add_page(page)

        with open(output_path, "wb") as output_file:
            pdf_write.write(output_file)

        return output_path
